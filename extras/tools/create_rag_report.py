from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path("/Users/atomas/Documents/TFM/entregables/propuesta_rag_incendios_espana.docx")

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "203548"
MUTED = "5F6B76"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "D9E0E7"
CALLOUT = "EEF5FB"
WHITE = "FFFFFF"
GOLD = "B07A14"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        prevent_row_split(row)
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths_dxa[idx] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_run_font(run, name="Calibri", size=11, bold=None, italic=None, color=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_hyperlink(paragraph, text, url, color=BLUE):
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")
    r_pr.append(r_fonts)
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    r_pr.append(c)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    r_pr.append(u)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "22")
    r_pr.append(sz)
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    set_run_font(run, size=9, color=MUTED)


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.1
    p.add_run(text)
    return p


def create_numbering_instance(doc):
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [
        int(node.get(qn("w:numId")))
        for node in numbering.findall(qn("w:num"))
    ]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1.")
    lvl.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    lvl.append(lvl_jc)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    p_pr.append(ind)
    lvl.append(p_pr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abs_ref = OxmlElement("w:abstractNumId")
    abs_ref.set(qn("w:val"), str(abstract_id))
    num.append(abs_ref)
    numbering.append(num)
    return num_id


def add_numbered(doc, text, num_id):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.1
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_pr.append(ilvl)
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.append(num_id_node)
    p_pr.append(num_pr)
    p.add_run(text)
    return p


def add_source_paragraph(doc, label, url, note):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f"{label}: ")
    set_run_font(r, bold=True, color=INK)
    add_hyperlink(p, url, url)
    if note:
        r2 = p.add_run(f". {note}")
        set_run_font(r2)


def add_callout(doc, title, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.left_indent = Inches(0.16)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.line_spacing = 1.1
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), CALLOUT)
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), BLUE)
    borders.append(left)
    p_pr.append(borders)
    r1 = p.add_run(f"{title}. ")
    set_run_font(r1, bold=True, color=DARK_BLUE)
    r2 = p.add_run(text)
    set_run_font(r2, color=INK)
    return p


def add_table(doc, headers, rows, widths_dxa, font_size=9):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = 0
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for idx, header in enumerate(headers):
        cell = hdr.cells[idx]
        set_cell_shading(cell, LIGHT_GRAY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(header)
        set_run_font(run, size=font_size, bold=True, color=INK)
    for row_data in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_data):
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx in (2, 3) else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(value)
            set_run_font(run, size=font_size, color=INK)
    set_table_geometry(table, widths_dxa)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Bullet 2", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(5)
        style.paragraph_format.line_spacing = 1.1


def build_document():
    doc = Document()
    configure_styles(doc)

    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    hr = hp.add_run("TFM · Sistema RAG para incendios y meteorología")
    set_run_font(hr, size=9, color=MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_after = Pt(0)
    fr = fp.add_run("Página ")
    set_run_font(fr, size=9, color=MUTED)
    add_page_field(fp)

    # Cover: editorial_cover pattern resolved for a technical report.
    doc.add_paragraph().paragraph_format.space_after = Pt(66)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(14)
    kr = kicker.add_run("PROPUESTA TÉCNICA")
    set_run_font(kr, size=11, bold=True, color=GOLD)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    tr = title.add_run("Sistema RAG de avisos meteorológicos\ne incendios forestales en España")
    set_run_font(tr, size=27, bold=True, color=INK)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(28)
    sr = subtitle.add_run(
        "Fuentes de información, conexiones, arquitectura de ingesta,\n"
        "estrategia de scraping y plan de implementación"
    )
    set_run_font(sr, size=14, color=DARK_BLUE)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_before = Pt(40)
    meta.paragraph_format.space_after = Pt(4)
    mr = meta.add_run("Trabajo Fin de Máster · Aplicación de detección de incendios")
    set_run_font(mr, size=11, bold=True, color=MUTED)
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = date_p.add_run("6 de julio de 2026")
    set_run_font(dr, size=11, color=MUTED)
    doc.add_page_break()

    doc.add_heading("1. Resumen ejecutivo", level=1)
    doc.add_paragraph(
        "El objetivo es incorporar a la aplicación de detección de incendios un sistema de "
        "generación aumentada por recuperación (RAG) capaz de responder con contexto reciente "
        "sobre avisos meteorológicos, riesgo de incendio y evolución de incendios forestales en España. "
        "El sistema no sustituirá a los organismos de emergencias: actuará como una capa informativa "
        "trazable que enlaza cada afirmación con su fuente y momento de actualización."
    )
    add_callout(
        doc,
        "Recomendación principal",
        "Construir una ingesta híbrida. Los avisos oficiales se consultan mediante API, CAP, RSS, "
        "CSV o servicios GIS; las noticias se descubren por RSS y se extraen de forma incremental; "
        "los foros se mantienen en un índice separado de baja confianza. Un batch exclusivamente "
        "diario no es suficiente para información operativa.",
    )
    add_bullet(doc, "Priorizar API y formatos estructurados antes que scraping.")
    add_bullet(doc, "Separar documentos —noticias, avisos y comunicados— de los eventos reales.")
    add_bullet(doc, "Aplicar filtros geográficos, temporales y de confianza en cada consulta.")
    add_bullet(doc, "Responder siempre con URL, fuente y fecha de actualización.")

    doc.add_heading("2. Contexto, alcance y objetivos", level=1)
    doc.add_paragraph(
        "El proyecto se desarrolla en el contexto de un Trabajo Fin de Máster centrado en una "
        "aplicación de detección de incendios. El RAG propuesto complementa la detección técnica "
        "con información pública reciente y explicable: qué avisos están vigentes, qué condiciones "
        "meteorológicas pueden favorecer la propagación, qué administraciones han comunicado un "
        "incendio y cómo ha evolucionado su estado."
    )
    doc.add_heading("2.1 Alcance funcional", level=2)
    add_bullet(doc, "Cobertura exclusiva de España, incluidas Canarias, Baleares, Ceuta y Melilla.")
    add_bullet(doc, "Contenido en castellano, catalán, gallego y euskera.")
    add_bullet(doc, "Avisos meteorológicos, riesgo de incendio, incendios activos e información contextual.")
    add_bullet(doc, "Consultas actuales e históricas con filtros por municipio, provincia y comunidad autónoma.")
    doc.add_heading("2.2 Límites", level=2)
    add_bullet(doc, "Una anomalía térmica satelital no se tratará automáticamente como incendio confirmado.")
    add_bullet(doc, "Los foros y redes sociales no determinarán el estado operativo de un incendio.")
    add_bullet(doc, "El RAG no emitirá instrucciones de emergencia ni sustituirá al 112.")

    doc.add_heading("3. Principios de diseño", level=1)
    principles_num_id = create_numbering_instance(doc)
    for text in (
        "Fuente estructurada primero: API REST, CAP/RSS, CSV, XML y WMS antes que HTML.",
        "Incrementalidad: descargar solamente URLs o recursos nuevos o modificados.",
        "Trazabilidad: conservar el original, la fecha de captura y la URL.",
        "Temporalidad explícita: distinguir fecha de publicación, actualización y periodo de validez.",
        "Confianza por fuente: priorizar organismos oficiales frente a prensa, foros o redes.",
        "Degradación controlada: si una fuente falla, conservar el último dato válido indicando su antigüedad.",
    ):
        add_numbered(doc, text, principles_num_id)

    doc.add_heading("4. Inventario de fuentes y tipo de conexión", level=1)
    doc.add_paragraph(
        "Las siguientes tablas resumen la conexión recomendada. “Gratuita” se refiere al acceso a la "
        "fuente; no incluye infraestructura, embeddings, modelos de lenguaje ni posibles permisos de "
        "redistribución."
    )

    doc.add_heading("4.1 Fuentes oficiales nacionales", level=2)
    add_table(
        doc,
        ["Fuente", "Conexión y formato", "Autenticación", "Coste", "Uso recomendado"],
        [
            ("AEMET: avisos", "RSS/Atom + CAP 1.2 (XML)", "No", "Gratis", "Avisos vigentes; 5–10 min"),
            ("AEMET OpenData", "API REST; JSON y descargas", "API key", "Gratis", "Histórico y backfill"),
            ("MITECO actuaciones", "PDF en URL estable", "No", "Gratis", "Parte provisional/definitivo"),
            ("MITECO EGIF", "Exportación XML/Excel", "No", "Gratis", "Histórico consolidado"),
            ("EFFIS/Copernicus", "WMS y descargas GIS", "Habitualmente no", "Gratis", "Riesgo, focos y perímetros"),
        ],
        [1500, 2300, 1500, 1150, 2910],
        font_size=8.5,
    )

    doc.add_heading("4.2 Fuentes autonómicas y meteorológicas", level=2)
    add_table(
        doc,
        ["Fuente", "Conexión", "Autenticación", "Coste", "Decisión técnica"],
        [
            ("Castilla y León", "API Opendatasoft; JSON/CSV/GeoJSON", "No", "Gratis", "Integrar por API"),
            ("Cataluña Bombers", "API Socrata/SODA", "Token opcional", "Gratis", "Integrar por API"),
            ("INFOCA Andalucía", "Dashboard ArcGIS", "Depende de capa", "Consulta gratis", "ArcGIS REST; HTML de respaldo"),
            ("Xunta de Galicia", "Comunicados HTML", "No", "Consulta gratis", "Scraping incremental"),
            ("Canarias", "Web dinámica/Ajax", "No", "Consulta gratis", "Scraping controlado"),
            ("GVA 112", "HTML y PDF", "No", "Consulta gratis", "No depender de API interna"),
            ("INFOAR Aragón", "HTML/PDF", "No", "Consulta gratis", "Scraping de comunicados"),
            ("Meteocat", "API REST JSON", "API key", "Gratis TFM", "Opcional; revisar difusión"),
            ("Euskalmet", "API REST JSON + JWT", "Registro gratuito", "Gratis", "API para alertas regionales"),
            ("MeteoGalicia", "RSS/XML y JSON", "No", "Gratis", "Feed/API; no scraping"),
        ],
        [1450, 2200, 1580, 1230, 2900],
        font_size=8.2,
    )

    doc.add_heading("4.3 Noticias y comunidad", level=2)
    add_table(
        doc,
        ["Fuente", "Conexión", "Acceso", "Ingesta", "Confianza"],
        [
            ("RTVE", "RSS general + páginas temáticas", "Gratuito", "RSS y extracción HTML", "Media-alta"),
            ("EFEverde", "RSS/HTML", "Web pública", "RSS + extracción limitada", "Media-alta"),
            ("Europa Press", "Web temática; feeds profesionales", "Web pública / feed de pago", "Metadatos + HTML limitado", "Media-alta"),
            ("Prensa regional", "RSS, sitemap o HTML", "Variable", "Por lista blanca", "Media"),
            ("Foros meteorológicos", "HTML; RSS si existe", "Gratuito", "Cada 2–6 horas", "Baja"),
            ("Reddit", "API OAuth", "Sujeto a condiciones", "Sólo si aporta valor", "Baja"),
        ],
        [1600, 2300, 1500, 2450, 1510],
        font_size=8.5,
    )

    doc.add_heading("5. Estrategia de conexión por tecnología", level=1)
    doc.add_heading("5.1 API REST y formatos estructurados", level=2)
    doc.add_paragraph(
        "Las API REST se consumirán con un cliente HTTP de servidor, reintentos exponenciales, timeout, "
        "caché y paginación. Para AEMET, Meteocat y Euskalmet las credenciales se almacenarán como "
        "secretos y nunca en el repositorio. Opendatasoft y Socrata permiten consultas JSON filtradas, "
        "evitando descargar datasets completos en cada ejecución."
    )
    doc.add_heading("5.2 CAP, RSS y Atom", level=2)
    doc.add_paragraph(
        "Los feeds funcionan como índices incrementales. El GUID, la URL y la fecha de actualización "
        "permiten detectar novedades. Los mensajes CAP de AEMET se procesarán como alertas estructuradas, "
        "no como texto periodístico."
    )
    doc.add_heading("5.3 WMS y datos GIS", level=2)
    doc.add_paragraph(
        "EFFIS se integrará mediante servicios WMS y descargas geoespaciales. Las consultas se limitarán "
        "al polígono de España y sus territorios insulares. Los focos satelitales se etiquetarán como "
        "detecciones o anomalías, manteniéndolos separados de incendios confirmados."
    )
    doc.add_heading("5.4 PDF y scraping HTML", level=2)
    doc.add_paragraph(
        "Los PDF de MITECO se descargarán sólo cuando cambie su hash o cabecera Last-Modified. Para HTML "
        "se utilizará extracción estática; Playwright quedará reservado para páginas JavaScript que no "
        "expongan una llamada JSON reutilizable y cuya automatización esté permitida."
    )

    doc.add_heading("6. Arquitectura propuesta", level=1)
    architecture_num_id = create_numbering_instance(doc)
    add_numbered(doc, "Descubrimiento: consultar API, feeds, sitemaps, listados y recursos conocidos.", architecture_num_id)
    add_numbered(doc, "Captura: descargar el recurso y conservar una copia original con fecha y hash.", architecture_num_id)
    add_numbered(doc, "Normalización: limpiar texto, convertir fechas, identificar idioma y fuente.", architecture_num_id)
    add_numbered(doc, "Filtrado: clasificar relevancia temática y comprobar que el evento afecta a España.", architecture_num_id)
    add_numbered(doc, "Enriquecimiento: extraer municipio, provincia, estado, severidad, superficie y periodo.", architecture_num_id)
    add_numbered(doc, "Resolución de eventos: asociar varias noticias o avisos a un mismo incendio.", architecture_num_id)
    add_numbered(doc, "Indexación: generar chunks y embeddings; actualizar búsqueda textual y vectorial.", architecture_num_id)
    add_numbered(doc, "Consulta RAG: recuperar con filtros, rerankear y responder con citas y fecha.", architecture_num_id)

    doc.add_heading("6.1 Componentes", level=2)
    add_bullet(doc, "Python: httpx, feedparser, lxml, trafilatura, pdfplumber/PyMuPDF y geopandas.")
    add_bullet(doc, "PostgreSQL: documentos, ejecuciones, estado de fuentes y metadatos.")
    add_bullet(doc, "PostGIS: geometrías, municipios y comprobación de pertenencia a España.")
    add_bullet(doc, "pgvector: embeddings en la misma base de datos.")
    add_bullet(doc, "Almacenamiento de originales: sistema de archivos o S3/MinIO.")
    add_bullet(doc, "Planificador: cron, APScheduler o Prefect según complejidad.")

    doc.add_heading("7. Modelo de información", level=1)
    doc.add_paragraph(
        "Se deben separar los documentos que describen la realidad de los eventos que representan esa "
        "realidad. Un incendio puede tener decenas de actualizaciones contradictorias; conservarlas como "
        "documentos independientes y resolverlas sobre un evento común evita perder la evolución."
    )
    add_table(
        doc,
        ["Entidad", "Campos esenciales", "Finalidad"],
        [
            (
                "Documento",
                "source, URL, title, body, published_at, fetched_at, updated_at, hash, language, trust",
                "Unidad de evidencia y citación",
            ),
            (
                "Evento",
                "type, name, municipality, province, geometry, start/end, status, severity, area_ha",
                "Representar incendio o episodio meteorológico",
            ),
            (
                "Observación",
                "event_id, document_id, field, value, observed_at, confidence",
                "Mantener versiones y discrepancias",
            ),
            (
                "Ejecución",
                "source_id, started_at, status, HTTP code, records, error, latency",
                "Monitorizar salud y frescura",
            ),
        ],
        [1300, 5200, 2860],
        font_size=9,
    )

    doc.add_heading("8. Frecuencia de actualización", level=1)
    add_table(
        doc,
        ["Canal", "Frecuencia sugerida", "Motivo"],
        [
            ("AEMET CAP", "5–10 minutos", "Los avisos pueden cambiar durante el episodio"),
            ("Incendios autonómicos", "10–30 minutos en campaña", "Fuente operativa más próxima"),
            ("MITECO", "30–60 minutos", "Parte provisional y consolidación diaria"),
            ("RSS/noticias", "20–30 minutos", "Equilibrio entre frescura y carga"),
            ("Foros", "2–6 horas", "Señal secundaria y no operativa"),
            ("Reconciliación", "Cada noche", "Reintentos, expiración, duplicados y reindexado"),
            ("EGIF/histórico", "Mensual o bajo demanda", "Datos consolidados y poco volátiles"),
        ],
        [2200, 2350, 4810],
        font_size=9,
    )
    add_callout(
        doc,
        "Decisión",
        "La arquitectura será de microbatches por fuente más un batch nocturno. Ejecutar todo una vez al "
        "día produciría respuestas desactualizadas precisamente cuando el riesgo sea mayor.",
    )

    doc.add_heading("9. Filtrado, deduplicación y confianza", level=1)
    doc.add_heading("9.1 Filtro geográfico", level=2)
    add_bullet(doc, "Reconocer municipios y provincias mediante nomenclátor oficial.")
    add_bullet(doc, "Geocodificar y comprobar intersección con el territorio español.")
    add_bullet(doc, "No aceptar una noticia como española únicamente por el dominio del medio.")
    add_bullet(doc, "Permitir sucesos fronterizos sólo cuando afecten territorio español.")
    doc.add_heading("9.2 Deduplicación", level=2)
    add_bullet(doc, "URL canónica y GUID de feed.")
    add_bullet(doc, "Hash del texto limpio.")
    add_bullet(doc, "SimHash/MinHash para republicaciones de agencia.")
    add_bullet(doc, "Clustering por municipio, fecha, distancia y tipo de evento.")
    doc.add_heading("9.3 Jerarquía de fuentes", level=2)
    trust_num_id = create_numbering_instance(doc)
    add_numbered(doc, "Organismos oficiales nacionales y autonómicos.", trust_num_id)
    add_numbered(doc, "Servicios públicos, RTVE y agencias informativas.", trust_num_id)
    add_numbered(doc, "Prensa nacional y regional incluida en lista blanca.", trust_num_id)
    add_numbered(doc, "Foros y redes sociales, siempre como señal comunitaria.", trust_num_id)

    doc.add_heading("10. Diseño del RAG", level=1)
    doc.add_paragraph(
        "La recuperación será híbrida: búsqueda textual BM25, embeddings multilingües y filtros de "
        "metadatos. Un reranker combinará similitud, confianza, proximidad geográfica y frescura."
    )
    add_bullet(doc, "Los avisos expirados se excluyen por defecto de preguntas sobre la situación actual.")
    add_bullet(doc, "Para “estado actual”, se recuperará la última observación oficial del evento.")
    add_bullet(doc, "Las contradicciones relevantes se mostrarán con sus respectivas horas y fuentes.")
    add_bullet(doc, "Cada respuesta incluirá enlaces y marcas temporales.")
    add_bullet(doc, "El contenido scrapeado se tratará como entrada no confiable frente a prompt injection.")

    doc.add_heading("11. Costes, licencias y cumplimiento", level=1)
    doc.add_paragraph(
        "AEMET, MITECO, EFFIS, Opendatasoft, Socrata y la mayoría de los portales públicos pueden "
        "utilizarse sin coste de acceso, aunque deben respetarse sus condiciones de atribución. Meteocat "
        "es gratuito para estudiantes e investigación previa solicitud; sus planes profesionales son de "
        "pago. Euskalmet requiere registro gratuito."
    )
    add_callout(
        doc,
        "Precaución jurídica",
        "Una página gratuita no implica una licencia para copiar y redistribuir íntegramente sus noticias. "
        "Cuando la licencia no sea clara se conservarán metadatos, URL, un extracto breve y hechos "
        "estructurados. No se sortearán paywalls, autenticación, CAPTCHA ni bloqueos.",
    )
    add_bullet(doc, "Mantener una ficha de licencia, atribución y frecuencia permitida por fuente.")
    add_bullet(doc, "Respetar robots.txt y condiciones de servicio.")
    add_bullet(doc, "Usar un User-Agent identificable y límites de petición conservadores.")
    add_bullet(doc, "Minimizar nombres de usuario y datos personales procedentes de foros.")
    add_bullet(doc, "Revisar permisos antes de una explotación pública o comercial.")

    doc.add_heading("12. Plan de implementación", level=1)
    add_table(
        doc,
        ["Fase", "Duración", "Entregables"],
        [
            ("1. Base", "Semana 1", "Esquema, registro de fuentes, AEMET CAP y almacenamiento original"),
            ("2. Oficial", "Semana 2", "AEMET OpenData, MITECO, EFFIS y Castilla y León"),
            ("3. Regional/noticias", "Semana 3", "Cataluña, Andalucía, RTVE y EFEverde"),
            ("4. Enriquecimiento", "Semana 4", "Geolocalización, eventos y deduplicación"),
            ("5. RAG", "Semana 5", "pgvector, recuperación híbrida y respuestas con citas"),
            ("6. Evaluación", "Semana 6", "Pruebas, observabilidad, documentación y memoria"),
        ],
        [1800, 1450, 6110],
        font_size=9,
    )

    doc.add_heading("12.1 MVP recomendado", level=2)
    add_bullet(doc, "AEMET CAP/RSS y AEMET OpenData.")
    add_bullet(doc, "MITECO: parte provisional y definitivo.")
    add_bullet(doc, "EFFIS WMS con filtro territorial español.")
    add_bullet(doc, "Castilla y León mediante Opendatasoft.")
    add_bullet(doc, "Cataluña mediante Socrata.")
    add_bullet(doc, "RTVE y EFEverde como contexto periodístico.")
    doc.add_paragraph(
        "Este alcance proporciona cobertura nacional, dos integraciones autonómicas sólidas y contexto "
        "informativo sin convertir el TFM en el mantenimiento simultáneo de diecisiete scrapers."
    )

    doc.add_heading("13. Evaluación", level=1)
    add_bullet(doc, "Frescura: minutos entre publicación e indexación.")
    add_bullet(doc, "Recall@k y nDCG de recuperación.")
    add_bullet(doc, "Precisión de municipio, estado, superficie y severidad extraídos.")
    add_bullet(doc, "Porcentaje de respuestas con citas que realmente respaldan la afirmación.")
    add_bullet(doc, "Tasa de duplicados y errores por fuente.")
    add_bullet(doc, "Pruebas temporales: activo frente a controlado, vigente frente a expirado.")
    doc.add_paragraph(
        "Se recomienda construir un conjunto de evaluación con entre 40 y 60 preguntas y una muestra "
        "manual de al menos 100 documentos/eventos. Debe incluir casos con actualizaciones sucesivas y "
        "fuentes contradictorias."
    )

    doc.add_heading("14. Conclusión", level=1)
    doc.add_paragraph(
        "El RAG es viable con un MVP prácticamente sin costes de acceso a fuentes. La decisión técnica "
        "clave es reservar el scraping para los lugares donde no exista una interfaz estructurada y "
        "mantener separados los niveles de confianza. La combinación de AEMET, MITECO, EFFIS, datos "
        "autonómicos y noticias seleccionadas permite construir un sistema reciente, trazable y adecuado "
        "para una aplicación de detección de incendios orientada a España."
    )

    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    doc.add_heading("Anexo A. Fuentes y documentación", level=1)
    sources = [
        ("AEMET — avisos CAP/RSS", "https://www.aemet.es/es/rss_info/avisos/esp", "Avisos vigentes por territorio."),
        ("AEMET OpenData", "https://opendata.aemet.es/centrodedescargas/AEMETApi", "API REST y documentación."),
        ("MITECO — actuaciones diarias", "https://www.miteco.gob.es/es/biodiversidad/temas/incendios-forestales/estadisticas-actuaciones.html", "Partes PDF y riesgo."),
        ("MITECO — EGIF", "https://www.miteco.gob.es/es/biodiversidad/temas/incendios-forestales/estadisticas-datos.html", "Histórico nacional."),
        ("EFFIS — datos y servicios", "https://forest-fire.emergency.copernicus.eu/applications/data-and-services", "WMS y descargas."),
        ("EFFIS — licencia", "https://forest-fire.emergency.copernicus.eu/about-effis/data-license", "Condiciones de reutilización."),
        ("Castilla y León — incendios", "https://jcyl.opendatasoft.com/explore/dataset/incendios-forestales/", "Opendatasoft."),
        ("Cataluña — actuaciones de Bombers", "https://interior.gencat.cat/ca/arees_dactuacio/bombers/actuacions-de-bombers/index.html", "Visor y datos abiertos."),
        ("Andalucía — INFOCA", "https://www.juntadeandalucia.es/organismos/ema/areas/incendios-forestales/situacion/incendios-activos.html", "Dashboard oficial."),
        ("Meteocat API", "https://apidocs.meteocat.gencat.cat/", "REST JSON con API key."),
        ("Euskalmet API", "https://opendata.euskadi.eus/api-euskalmet/", "REST JSON con registro."),
        ("MeteoGalicia — avisos", "https://www.meteogalicia.gal/web/predicion/avisos?request_locale=es", "RSS y documentación."),
        ("RTVE RSS", "https://www.rtve.es/rss/", "Canales oficiales."),
        ("RTVE — incendios", "https://www.rtve.es/temas/incendios/3690/", "Página temática."),
        ("EFEverde — incendios", "https://efeverde.com/desastres/incendios/", "Información ambiental."),
        ("Europa Press — incendios", "https://www.europapress.es/temas/incendios-forestales/", "Página temática."),
        ("Cazatormentas", "https://cazatormentas.com/foro/index.php", "Foro comunitario."),
        ("Foro Tiempo.com", "https://foro.tiempo.com/meteorologia-general-b1.0/", "Foro meteorológico."),
    ]
    for label, url, note in sources:
        if label == "Meteocat API":
            doc.add_heading("Fuentes meteorológicas, informativas y comunitarias", level=2)
        add_source_paragraph(doc, label, url, note)

    doc.add_heading("Anexo B. Checklist de puesta en marcha", level=1)
    for item in (
        "Registrar cada fuente con propietario, método de conexión, frecuencia, licencia y nivel de confianza.",
        "Guardar las API keys en secretos de servidor y excluirlas del control de versiones.",
        "Conservar una copia original, hash y fecha de captura antes de transformar cada recurso.",
        "Implementar ETag, Last-Modified, caché, timeout, reintentos y límites de petición.",
        "Validar el filtro territorial con geometrías de España, islas, Ceuta y Melilla.",
        "Separar incidentes confirmados, anomalías térmicas, noticias y contenido comunitario.",
        "Comprobar que los avisos expirados no aparecen como vigentes.",
        "Probar cambios de estado: activo, estabilizado, controlado y extinguido.",
        "Monitorizar retraso de actualización, errores, volumen y cambios de estructura por fuente.",
        "Exigir fuente, URL y marca temporal en toda respuesta generada por el RAG.",
    ):
        add_bullet(doc, item)

    props = doc.core_properties
    props.title = "Propuesta técnica para un RAG de avisos meteorológicos e incendios en España"
    props.subject = "Trabajo Fin de Máster"
    props.author = "Equipo del TFM"
    props.keywords = "RAG, incendios forestales, AEMET, scraping, España, meteorología"
    props.comments = "Documento técnico generado para la planificación del TFM."

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build_document())
